# ================================================================
# HVIP TELEMATICS REPORT GENERATOR
# ================================================================

# This script analyzes vehicle GPS data inside priority population zones
# and generates [3]:
#   - CSV report of time and miles in each zone per vehicle
#   - Map image of vehicle activity overlaying priority zones
#   - Word report combining the table and the map (landscape)
#
# -----------------------------
# Step-by-Step Preparatory Instructions
# -----------------------------
#
# 1 Install Python 3
#     - Go to https://www.python.org/downloads/
#     - Download and install Python 3
#     - On Windows, check "Add Python to PATH"
#
# 2 Prepare input files
#     - Create a folder called "Telematics"
#     - Save this script file inside your new Telematics folder
#     - Copy the Priority Population 4.0 Geodatabase you downloaded from the readme file link into the Telematics folder
#     - Place your Telematics CSV in the Telematics Folder
#     - Your Telematics CSV MUST include headers EXACTLY matching: latitude, longitude, mileage, orig_time, license_nmbr
#
# 3 Install required packages (if not already installed)
#     - Open Command Prompt or Terminal
#     - Navigate to the folder where this script is saved
#     - Run each command:
#           pip install geopandas
#           pip install fiona
#           pip install pandas
#           pip install matplotlib
#           pip install contextily
#           pip install python-docx
#           pip install shapely
#
# Outputs for this report will be saved in an "outputs" folder automatically

# ================================================================
## BEGIN CODE - import packages utilized throughoput the code
# ================================================================

from pathlib import Path
import geopandas as gpd
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from shapely.geometry import box
from mpl_toolkits.axes_grid1.inset_locator import inset_axes
from datetime import datetime
import warnings
import fiona

warnings.filterwarnings("ignore")

# -----------------------------------------------------------------
# RUN DATE - To later include the date in the report automatically
# -----------------------------------------------------------------
RUN_DATE = datetime.now().strftime("%Y%m%d")
RUN_DATE_HUMAN = datetime.now().strftime("%Y-%m-%d")

# ----------------------------------------------------------
# INPUT USER SETTINGS - ADD YOUR PATHS HERE WHERE INDICATED [4]
# ----------------------------------------------------------
COMPANY_NAME = "CALSTART TEST COMPANY" #<<<<<<<<<<<<<<<<<---YOUR COMPANY NAME
BASE_DIR = Path("C:/Users/mbrown/Desktop/Telematics")#<<<<<<<<<<<<<<<<<---Your root folder

GDB_PATH = BASE_DIR / "Priority Populations 4.0 Geodatabase" / "Priority Populations 4.0 Combined Layer.gdb"#<<<<<<<<<<<<<<<<<---Your CCI .gdb full path
CSV_PATH = BASE_DIR / "CaliforniaHVIP_sampleData_AVTA_test.csv"#<<<<<<<<<<<<<<<<<--- Your Specific dataset with latitude, longitude, mileage, orig_time, license_nmbr

# ---------------------------------------------------------------------------------------------
# OUTPUT USER SETTINGS - this is what your output files and paths will be named automatically
# ---------------------------------------------------------------------------------------------

def make_folder(path: Path):
    path.mkdir(parents=True, exist_ok=True)
    return path

# Base output directories
OUTPUT_DIR = make_folder(BASE_DIR / "outputs")
COMBINED_FOLDER = make_folder(OUTPUT_DIR / "combined_priority_population")

# Output file paths
OUTPUT_COMBINED_PATH = COMBINED_FOLDER / "combined_priority_population.shp"
OUTPUT_CALIFORNIA_BOUNDARY = COMBINED_FOLDER / "california_boundary.shp"

OUTPUT_FINAL_REPORT = OUTPUT_DIR / f"{COMPANY_NAME}_telematicsreport_{RUN_DATE}.csv"
MAP_PATH = OUTPUT_DIR / "priority_zone_map.png"
WORD_REPORT_PATH = OUTPUT_DIR / f"{COMPANY_NAME}_telematicsreport_{RUN_DATE}.docx"

# ----------------------------------------------------------------------------------
# HELPER FUNCTIONS - will remove/tidy temporary columns/formatting at later stages
# ----------------------------------------------------------------------------------
def drop_index_right(gdf):
    """Drop 'index_right' column if it exists."""
    return gdf.drop(columns=["index_right"], errors='ignore')

def remove_spacing(paragraph):
    """Remove spacing before and after a paragraph in Word."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), "0")
    spacing.set(qn('w:after'), "0")
    pPr.append(spacing)

# ================================================================
# STEP 1: LOAD & COMBINE PRIORITY POPULATION ZONES
# ================================================================
print("Loading priority population geodatabase...")

layers = fiona.listlayers(GDB_PATH)

# Layers to exclude for priority zones only
EXCLUDE_LAYERS = [
    "DAC_half_mile_neighbor__low_income_household_eligible",
    "Low_income_household_eligible",
    "Not_a_priority_population_area_low_income_households_are_eligible"
]

# Load all layers once
all_gdfs = []
priority_gdfs = []

for layer in layers:
    gdf = gpd.read_file(GDB_PATH, layer=layer)
    gdf["SourceLayer"] = layer
    all_gdfs.append(gdf)
    if layer not in EXCLUDE_LAYERS:
        priority_gdfs.append(gdf)

# Combined priority zones
combined_gdf = gpd.GeoDataFrame(pd.concat(priority_gdfs, ignore_index=True))
combined_gdf.to_file(OUTPUT_COMBINED_PATH)
print(f"Saved combined priority zones to: {OUTPUT_COMBINED_PATH}")

# California boundary
all_layers_gdf = gpd.GeoDataFrame(pd.concat(all_gdfs, ignore_index=True))
all_layers_gdf = drop_index_right(all_layers_gdf)
all_layers_gdf.to_file(OUTPUT_CALIFORNIA_BOUNDARY)
print(f"California boundary saved to: {OUTPUT_CALIFORNIA_BOUNDARY}")

# ================================================================
# STEP 2: LOAD TELEMATICS CSV
# ================================================================
print("Loading telematics data...")
telematics_df = pd.read_csv(CSV_PATH)
telematics_gdf = gpd.GeoDataFrame(
    telematics_df,
    geometry=gpd.points_from_xy(telematics_df['longitude'], telematics_df['latitude']),
    crs="EPSG:4326"
).to_crs(all_layers_gdf.crs)

# ================================================================
# STEP 3: CATEGORIZE POINTS AND CALCULATE PER-VEHICLE METRICS
# ================================================================
print("Categorizing points and calculating per-vehicle metrics...")

# Ensure datetime and sort
telematics_gdf["orig_time"] = pd.to_datetime(telematics_gdf["orig_time"])
telematics_gdf = telematics_gdf.sort_values(by=["license_nmbr", "orig_time"])

# Calculate time difference in hours and mileage difference
telematics_gdf["time_diff_hr"] = telematics_gdf.groupby("license_nmbr")["orig_time"].diff().dt.total_seconds() / 3600
telematics_gdf["mileage_diff"] = telematics_gdf.groupby("license_nmbr")["mileage"].diff()

# A vehicle is only "active" during an interval if mileage actually changed.
# This prevents long idle/parked gaps (e.g. overnight) from being counted as
# drive hours just because two consecutive GPS pings are far apart in time.
telematics_gdf["is_active"] = telematics_gdf["mileage_diff"] > 0

# If ignition_status is present in the data, use it as an additional clue
# an interval only counts as driving if ignition was on AND mileage moved, otherwise its idle
if "ignition_status" in telematics_gdf.columns:
    telematics_gdf["is_active"] = telematics_gdf["is_active"] & (telematics_gdf["ignition_status"] == 1)

# Zero out time_diff_hr for any interval where the vehicle wasn't actually moving
telematics_gdf["time_diff_hr"] = telematics_gdf["time_diff_hr"].where(telematics_gdf["is_active"], 0)


# ---------------------------------------
# Step 3a: Points in priority populations
# ---------------------------------------
priority_join = gpd.sjoin(telematics_gdf, combined_gdf, how="left", predicate="intersects")
priority_join["Category"] = priority_join["index_right"].notnull().map({True: "Priority Populations", False: None})
priority_join = drop_index_right(priority_join)

# --------------------------------------------
# Step 3b: Points not in priority populations
# --------------------------------------------
not_priority = priority_join[priority_join["Category"].isna()].copy()
california_gdf_safe = drop_index_right(all_layers_gdf.copy())

california_join = gpd.sjoin(not_priority, california_gdf_safe, how="left", predicate="intersects")
california_join["Category"] = california_join["index_right"].notnull().map({True: "Inside California", False: "Outside California"})
california_join = drop_index_right(california_join)

# --------------------------------------------
# Step 3c: Combine all categorized points
# --------------------------------------------
final_gdf = pd.concat([
    priority_join[priority_join["Category"].notna()],
    california_join
], ignore_index=True)

# ---------------------------------------
# Step 3d: Aggregate metrics per vehicle
# ---------------------------------------
agg_fields = final_gdf.groupby(["license_nmbr", "Category"]).agg(
    Hours=("time_diff_hr", "sum"),
    Miles=("mileage_diff", "sum"),
    Days=("orig_time", lambda x: x.dt.date.nunique())
).reset_index()

hours_pivot = agg_fields.pivot_table(index="license_nmbr", columns="Category", values="Hours", fill_value=0)
miles_pivot = agg_fields.pivot_table(index="license_nmbr", columns="Category", values="Miles", fill_value=0)
days_pivot = agg_fields.pivot_table(index="license_nmbr", columns="Category", values="Days", fill_value=0)

for col in ["Priority Populations", "Inside California", "Outside California"]:
    for pivot in [hours_pivot, miles_pivot, days_pivot]:
        if col not in pivot.columns:
            pivot[col] = 0

# ----------------------------------
# Step 3e: Build per-vehicle csv report
# ----------------------------------
report = pd.DataFrame(index=hours_pivot.index)
report["Hours_Outside_CA"] = hours_pivot["Outside California"]
report["Hours_Priority_Populations"] = hours_pivot["Priority Populations"]
report["Hours_Inside_CA"] = hours_pivot["Inside California"]
report["Total_Hours"] = report[["Hours_Outside_CA", "Hours_Priority_Populations", "Hours_Inside_CA"]].sum(axis=1)
report["%_Hours_Outside_CA"] = (report["Hours_Outside_CA"] / report["Total_Hours"] * 100).round(2)
report["%_Hours_Priority_Populations"] = (report["Hours_Priority_Populations"] / report["Total_Hours"] * 100).round(2)

report["Miles_Outside_CA"] = miles_pivot["Outside California"]
report["Miles_Priority_Populations"] = miles_pivot["Priority Populations"]
report["Miles_Inside_CA"] = miles_pivot["Inside California"]
report["Total_Miles"] = report[["Miles_Outside_CA", "Miles_Priority_Populations", "Miles_Inside_CA"]].sum(axis=1)
report["%_Miles_Outside_CA"] = (report["Miles_Outside_CA"] / report["Total_Miles"] * 100).round(2)
report["%_Miles_Priority_Populations"] = (report["Miles_Priority_Populations"] / report["Total_Miles"] * 100).round(2)

report["%_Days_KeyOn_Priority"] = (days_pivot["Priority Populations"] / days_pivot[["Priority Populations", "Inside California", "Outside California"]].sum(axis=1) * 100).round(2)

report = report.reset_index().rename(columns={"license_nmbr": "Telematics"})

# -------------------
# Step 3f: Save CSV
# -------------------
report.to_csv(OUTPUT_FINAL_REPORT, index=False)
print(f"Final per-vehicle report saved to: {OUTPUT_FINAL_REPORT}")

# ==================================================
# STEP 4: GENERATE VEHICLE ACTIVITY MAP WITH INSET
# ==================================================
print("Generating vehicle activity map with inset...")

MAP_TITLE = "Vehicle Activity Within Priority Population Zones - beta testing"
BUFFER = 0.05

minx, miny, maxx, maxy = telematics_gdf.total_bounds
x_buffer = (maxx - minx) * BUFFER
y_buffer = (maxy - miny) * BUFFER

fig, ax = plt.subplots(figsize=(12, 12))
combined_gdf.plot(ax=ax, color="lightblue", edgecolor="black", alpha=0.5, linewidth=0.5)
telematics_gdf.plot(ax=ax, color="red", markersize=10, alpha=0.7)
ax.set_xlim(minx - x_buffer, maxx + x_buffer)
ax.set_ylim(miny - y_buffer, maxy + y_buffer)
ax.set_title(MAP_TITLE, fontsize=14)
ax.axis("off")

# Inset map: full California with extent rectangle
ax_inset = inset_axes(ax, width="30%", height="30%", loc='upper right')
all_layers_gdf.plot(ax=ax_inset, color="lightgrey", edgecolor="black")
extent_rect = box(minx, miny, maxx, maxy)
gpd.GeoSeries([extent_rect], crs=telematics_gdf.crs).plot(ax=ax_inset, facecolor="none", edgecolor="red", linewidth=2)
ax_inset.set_xticks([])
ax_inset.set_yticks([])
ax_inset.set_title("California Overview", fontsize=10)

plt.savefig(MAP_PATH, dpi=300, bbox_inches="tight")
plt.close()
print(f"Map saved to: {MAP_PATH}")

# ============================================================
# STEP 5: GENERATE EXECUTIVE-STYLE SUMMARY WORD REPORT (AKA ONE-PAGER)
# ============================================================
print("Generating one-page executive summary Word report...")

fleet_totals = {
    "Total Hours": report["Total_Hours"].sum(),
    "Hours in Priority Populations": report["Hours_Priority_Populations"].sum(),
    "% Hours in Priority Populations": round(report["Hours_Priority_Populations"].sum() / report["Total_Hours"].sum() * 100, 2),
    "Hours Outside CA": report["Hours_Outside_CA"].sum(),
    "Total Miles": report["Total_Miles"].sum(),
    "Miles in Priority Populations": report["Miles_Priority_Populations"].sum(),
    "% Miles in Priority Populations": round(report["Miles_Priority_Populations"].sum() / report["Total_Miles"].sum() * 100, 2),
    "Miles Outside CA": report["Miles_Outside_CA"].sum()
}

doc = Document()

# Main header
h1 = doc.add_heading("Fleetwide Telematics Summary", level=1)
remove_spacing(h1)

# Subheader: company name
h2 = doc.add_heading(COMPANY_NAME, level=2)
remove_spacing(h2)

# Sub-subheader: run date
h3 = doc.add_heading(RUN_DATE_HUMAN, level=3)

# Intro paragraph in italics, font size 12
p_intro = doc.add_paragraph(
    "This report summarizes the amount of time and distance each vehicle spent "
    "inside designated priority population areas. It is generated automatically "
    "from telematics data and GIS zone definitions. For vehicle-level details, "
    "please reference the supplemental Excel output."
)
p_intro.runs[0].italic = True
p_intro.runs[0].font.size = Pt(11)

# Fleet metrics summary
summary_text = (
    f"Fleet-level telematics data shows the following key metrics:\n\n"
    f"• Total Hours: {fleet_totals['Total Hours']:.1f}\n"
    f"• Hours in Priority Populations: {fleet_totals['Hours in Priority Populations']:.1f} "
    f"({fleet_totals['% Hours in Priority Populations']}%)\n"
    f"• Hours Outside California: {fleet_totals['Hours Outside CA']:.1f}\n"
    f"• Total Miles: {fleet_totals['Total Miles']:.1f}\n"
    f"• Miles in Priority Populations: {fleet_totals['Miles in Priority Populations']:.1f} "
    f"({fleet_totals['% Miles in Priority Populations']}%)\n"
    f"• Miles Outside California: {fleet_totals['Miles Outside CA']:.1f}\n"
)

p_metrics = doc.add_paragraph()
run_metrics = p_metrics.add_run(summary_text)
run_metrics.font.size = Pt(11)

# Vehicle activity map
doc.add_heading("Vehicle Activity Map", level=2)
doc.add_picture(str(MAP_PATH), width=Inches(3), height=Inches(4.5))

# Save document
doc.save(WORD_REPORT_PATH)
print(f"Executive summary saved to: {WORD_REPORT_PATH}")
print("All steps complete. Code execution finished successfully.")
